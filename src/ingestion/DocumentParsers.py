from abc import ABC, abstractmethod
from typing import Dict, Any, List
import pdfplumber
import docx
from io import BytesIO
from langsmith import traceable


class DocumentParser(ABC):
    """Abstract base class for document parsers"""
    
    @abstractmethod
    def parse(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Parse document and return text with metadata"""
        pass


class PDFParser(DocumentParser):
    """PDF document parser"""
    
    def parse(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Parse a PDF. Uses pdfplumber, which preserves word spacing far better than PyPDF2
        (PyPDF2 mangled runs into 'Implementedserver-side', hurting embeddings + readability)."""
        try:
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
                num_pages = len(pdf.pages)
            text = "\n".join(pages)

            return {
                'text': text.strip(),
                'metadata': {
                    'filename': filename,
                    'file_type': 'pdf',
                    'num_pages': num_pages
                }
            }
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {str(e)}")


class DOCXParser(DocumentParser):
    """DOCX document parser"""
    
    def parse(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            doc = docx.Document(BytesIO(file_bytes))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                'text': text.strip(),
                'metadata': {
                    'filename': filename,
                    'file_type': 'docx',
                    'num_paragraphs': len(doc.paragraphs)
                }
            }
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {str(e)}")


class TXTParser(DocumentParser):
    """Plain text document parser"""
    
    def parse(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Parse text document"""
        try:
            text = file_bytes.decode('utf-8')
            
            return {
                'text': text.strip(),
                'metadata': {
                    'filename': filename,
                    'file_type': 'txt',
                    'char_count': len(text)
                }
            }
        except Exception as e:
            raise RuntimeError(f"Failed to parse text file: {str(e)}")


class DocumentParserFactory:
    """Factory for creating appropriate document parsers"""
    
    _parsers = {
        '.pdf': PDFParser,
        '.docx': DOCXParser,
        '.txt': TXTParser,
        '.md': TXTParser
    }
    
    @classmethod
    def get_parser(cls, file_extension: str) -> DocumentParser:
        """Get appropriate parser for file extension"""
        parser_class = cls._parsers.get(file_extension.lower())
        if not parser_class:
            raise ValueError(f"Unsupported file type: {file_extension}")
        return parser_class()
    
    @classmethod
    def supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions"""
        return list(cls._parsers.keys())


@traceable(name="parse_document")
def parse_document(file_bytes: bytes, file_extension: str = None, filename: str = 'uploaded_file') -> Dict[str, Any]:
    """Convenience function to parse any supported document"""
    if file_extension is None:
        raise ValueError("file_extension must be provided when using bytes input")

    parser = DocumentParserFactory.get_parser(file_extension)
    result = parser.parse(file_bytes, filename)
    print(f"[PARSE] Parsed {filename}: {len(result['text'])} chars, metadata: {result['metadata']}")
    return result